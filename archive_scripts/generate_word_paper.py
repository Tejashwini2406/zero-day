#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
from pathlib import Path

tex_path = Path(__file__).resolve().parents[1] / 'paper' / 'zero_day_ieee.tex'
out_path = Path(__file__).resolve().parents[1] / 'paper' / 'zero_day_ieee.docx'

if not tex_path.exists():
    print(f"LaTeX source not found at {tex_path}")
    exit(1)

with open(tex_path, 'r', encoding='utf-8') as f:
    tex = f.read()

def extract_between(s, start_marker, end_marker):
    si = s.find(start_marker)
    if si == -1:
        return ''
    si += len(start_marker)
    ei = s.find(end_marker, si)
    if ei == -1:
        return s[si:].strip()
    return s[si:ei].strip()

import re
m = re.search(r"\\title\{(.+?)\}", tex, re.S)
title = m.group(1).strip() if m else 'Zero-Day PoC'

m = re.search(r"\\author\{(.+?)\}", tex, re.S)
author = m.group(1).strip() if m else 'Author'

abstract = extract_between(tex, '\\begin{abstract}', '\\end{abstract}')

sections = []
for sec in re.split(r"\\section\{", tex)[1:]:
    title_end = sec.find('}')
    sec_title = sec[:title_end].strip()
    content = sec[title_end+1:]
    content = re.split(r"\\section\{|\\end\{document\}", content)[0].strip()
    content = re.sub(r"\\\\[a-zA-Z]+\{.*?\}", '', content)
    content = re.sub(r"\\\\", '\n', content)
    content = re.sub(r"\$", '', content)
    sections.append((sec_title, content.strip()))

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

h = doc.add_heading(level=0)
h_run = h.add_run(title)
h_run.bold = True

doc.add_paragraph(author)

doc.add_heading('Abstract', level=1)
doc.add_paragraph(abstract)

for sec_title, sec_content in sections:
    doc.add_heading(sec_title, level=1)
    for p in re.split(r"\n\s*\n", sec_content):
        p = p.strip()
        if p:
            doc.add_paragraph(p)

bib = extract_between(tex, '\\begin{thebibliography}', '\\end{thebibliography}')
if bib:
    doc.add_heading('References', level=1)
    for line in bib.splitlines():
        line = line.strip()
        if line:
            line = re.sub(r"\\\\bibitem\{.*?\}", '', line)
            doc.add_paragraph(line)

out_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(out_path)
print(f"Wrote {out_path}")
