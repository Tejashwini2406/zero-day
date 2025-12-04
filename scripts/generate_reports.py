#!/usr/bin/env python3
"""Generate DOCX report files used for handoff/jury.
Creates three files under report/: 2. ABSTRACT ACKNOWLEDGEMANT.docx, 3. CHAPTERS.docx, Solving Complex Engineering Problems.docx
"""
from docx import Document
from pathlib import Path

OUT = Path('report')
OUT.mkdir(exist_ok=True)

def make_abstract():
    doc = Document()
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This report describes a proof-of-concept Zero-Day Detection Framework that ingests telemetry, constructs temporal graphs, trains and serves detection models, and provides explainable outputs mapped to MITRE ATT&CK techniques. The system targets Kubernetes environments and includes a Graph Builder, streaming ingestion, ML baselines, TGNN scaffolding, an inference pipeline and a containment operator.'
    )
    doc.add_heading('Acknowledgements', level=2)
    doc.add_paragraph('We thank the project contributors and test operators for validating the PoC in Minikube.')
    path = OUT / '2. ABSTRACT ACKNOWLEDGEMANT.docx'
    doc.save(path)
    print('Wrote', path)

def make_chapters():
    doc = Document()
    doc.add_heading('Chapters', level=1)
    chapters = [
        '1. Introduction',
        '2. Background and Related Work',
        '3. System Architecture',
        '  3.1 Telemetry Collection',
        '  3.2 Message Bus & Streaming',
        '  3.3 Temporal Graph Builder',
        '  3.4 ML Models and TGNN',
        '  3.5 Inference and Explainability',
        '  3.6 Containment and Automation',
        '4. Implementation',
        '5. Experiments and Demo',
        '6. Limitations and Future Work',
        '7. Conclusion',
        '8. Appendices (scripts, manifests, runbook)'
    ]
    for c in chapters:
        doc.add_paragraph(c)
    path = OUT / '3. CHAPTERS.docx'
    doc.save(path)
    print('Wrote', path)

def make_essay():
    doc = Document()
    doc.add_heading('Solving Complex Engineering Problems', level=1)
    doc.add_paragraph('This document outlines the engineering approach used to design and iterate on the Zero-Day Detection Framework. It covers problem decomposition, incremental prototyping, fault-tolerant choices, and strategies for delivering a robust demo while reducing blast radius.')
    doc.add_paragraph('Key lessons:')
    doc.add_paragraph('- Build minimal reproducible components first (file mode graph builder).')
    doc.add_paragraph('- Use fallbacks to keep end-to-end demoable under constrained environments.')
    doc.add_paragraph('- Separate heavy training workloads from runtime inference for operational safety.')
    path = OUT / 'Solving Complex Engineering Problems.docx'
    doc.save(path)
    print('Wrote', path)

def main():
    make_abstract()
    make_chapters()
    make_essay()

if __name__ == '__main__':
    main()
