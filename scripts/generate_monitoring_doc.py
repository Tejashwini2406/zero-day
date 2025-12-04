#!/usr/bin/env python3
from docx import Document
from pathlib import Path

OUT = Path('report')
OUT.mkdir(exist_ok=True)

def make_monitoring():
    doc = Document()
    doc.add_heading('Monitoring & Dashboard', level=1)
    doc.add_paragraph('This document summarizes the Prometheus + Grafana monitoring deployed for the PoC and includes the exported Grafana dashboard JSON used for provisioning.')
    jf = Path('report/grafana_dashboard.json')
    if jf.exists():
        doc.add_heading('Grafana Dashboard (exported JSON)', level=2)
        with jf.open() as fh:
            content = fh.read()
        # Insert as a code-like paragraph (docx has limited preformatting)
        para = doc.add_paragraph()
        para.style = 'Intense Quote'
        for line in content.splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph('Grafana dashboard JSON not found at report/grafana_dashboard.json')

    doc.add_heading('How to view', level=2)
    doc.add_paragraph('Port-forward Grafana and visit http://127.0.0.1:3000 (admin/admin). The dashboard "ZeroDay PoC Metrics" is provisioned automatically.')

    path = OUT / '4. MONITORING.docx'
    doc.save(path)
    print('Wrote', path)

if __name__ == '__main__':
    make_monitoring()
